from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Créer le document
doc = Document()
doc.core_properties.author = 'Olivier Lulihoshi Matabaro'
doc.core_properties.title = 'CV Humanitaire'

# Ajouter en-tête avec nom et contact
header = doc.sections[0].header
p_header = header.paragraphs[0]
p_header.text = 'LULIHOSHI MATABARO Olivier | Goma, RDC | +243 997421542 / +243 812532644 | Olivierrishimata@gmail.com'
p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Ajouter photo
doc.add_picture('/mnt/data/fc93ba80-ca08-4356-8bc2-ab527a1bc2ad.jpeg', width=Inches(1.5))

# Titre du CV
title = doc.add_paragraph('Curriculum Vitae')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(20)
title.runs[0].font.bold = True

# Exemple : Ajouter un tableau pour la formation
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
table.autofit = True
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Année'
hdr_cells[1].text = 'Établissement'
hdr_cells[2].text = 'Diplôme / Certification'

# Exemple de ligne de formation
row_cells = table.add_row().cells
row_cells[0].text = '2021-2023'
row_cells[1].text = 'ISDR Grands-Lacs, Goma'
row_cells[2].text = 'Licence (Bac+5) en Gestion des Entreprises de Développement'

# Ajouter une section Expérience (exemple)
doc.add_paragraph('Expériences Professionnelles', style='Heading 1')
doc.add_paragraph('Assistant Logistique (Bénévole) - FUPRODI Asbl – Goma, RDC | Juillet 2025 – Présent')
doc.add_paragraph('- Coordination et suivi des opérations logistiques pour les projets humanitaires.')
doc.add_paragraph('- Gestion des stocks, inventaires et distribution de matériel aux bénéficiaires.')
doc.add_paragraph('- Support à l’équipe dans la planification et l’organisation des missions sur le terrain.')
doc.add_paragraph('- Optimisation des processus pour améliorer l’efficacité des interventions.')

# Sauvegarder le document
file_path = '/mnt/data/CV_Olivier_Humanitaire.docx'
doc.save(file_path)
file_path
